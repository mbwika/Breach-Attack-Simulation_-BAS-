# speech_attack_sim/test_harness/run_pipeline.py
import sys
import os
import argparse
# Add the project root to Python path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import requests
import datetime
import importlib
from attacks.prompt_injection import generate_injected_prompt, INJECTION_PHRASES
from attacks.adversarial_audio import generate_adversarial_audio
from test_harness.logger import log_result, log_result_to_file
import json
import re
import tempfile
import time
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Argument parser for runtime configuration
parser = argparse.ArgumentParser(description="Run GenAI-ML attack Simulation pipeline against multiple API endpoints dynamically.")
parser.add_argument('--config', type=str, default="test_harness/api_config.json", help="Path to API config JSON file")
parser.add_argument('--input', type=str, help="Input text for attack simulation")
parser.add_argument('--email', type=str, help="Email address for endpoints that require it")
parser.add_argument('--file', type=str, help="Path to file for endpoints that require file upload")
parser.add_argument('--url', type=str, help="URL for endpoints that require it")
parser.add_argument('--country', type=str, help="Country for endpoints that require it")
args = parser.parse_args()

SAMPLE_TEXT = args.input
USER_EMAIL = args.email
USER_FILE = args.file
USER_URL = args.url
USER_COUNTRY = args.country
CONFIG_PATH = args.config

def normalize_email(email):
    # Basic normalization: lowercase and strip
    if not email:
        return None
    email = email.strip().lower()
    # Basic validation
    if re.match(r"^[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,}$", email):
        return email
    return None

def validate_url(url):
    # Basic URL validation
    if not url:
        return None
    url = url.strip()
    if re.match(r"^https?://[\w\.-]+(:\d+)?(/[\w\.-]*)*", url):
        return url
    return None

def generate_malicious_docx(malicious_text):
    """Generate a .docx file containing malicious text"""
    try:
        doc = Document()
        doc.add_heading('Resume', 0)
        doc.add_paragraph('Name: John Doe')
        doc.add_paragraph('Email: john.doe@example.com')
        doc.add_paragraph('Experience:')
        # Embed malicious text in the document
        doc.add_paragraph(f'Special Skills: {malicious_text}')
        doc.add_paragraph('Education: University of Example')
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        doc.save(temp_file.name)
        temp_file.close()
        return temp_file.name
    except ImportError:
        print("Warning: python-docx not available, creating dummy .docx file")
        temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False, mode='w')
        temp_file.write(f"Resume with malicious content: {malicious_text}")
        temp_file.close()
        return temp_file.name

def generate_malicious_pdf(malicious_text):
    """Generate a .pdf file containing malicious text"""
    try:
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        c = canvas.Canvas(temp_file.name, pagesize=letter)
        c.drawString(100, 750, "Resume")
        c.drawString(100, 720, "Name: John Doe")
        c.drawString(100, 690, "Email: john.doe@example.com")
        c.drawString(100, 660, "Experience:")
        # Embed malicious text in the PDF
        c.drawString(100, 630, f"Special Skills: {malicious_text[:80]}")  # Truncate for PDF
        c.drawString(100, 600, "Education: University of Example")
        c.save()
        temp_file.close()
        return temp_file.name
    except ImportError:
        print("Warning: reportlab not available, creating dummy .pdf file")
        temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False, mode='w')
        temp_file.write(f"Resume with malicious content: {malicious_text}")
        temp_file.close()
        return temp_file.name

def generate_malicious_json(malicious_text):
    """Generate a .json file containing malicious text"""
    temp_file = tempfile.NamedTemporaryFile(suffix='.json', delete=False, mode='w')
    malicious_data = {
        "profile": {
            "name": "John Doe",
            "email": "john.doe@example.com",
            "skills": malicious_text,
            "experience": "5 years in software development"
        }
    }
    json.dump(malicious_data, temp_file, indent=2)
    temp_file.close()
    return temp_file.name

def get_param_value(param, attack_values=None):
    # Get value for parameter from CLI args, attack values, config defaults, or prompt
    name = param.get('name')
    param_type = param.get('type')
    normalize = param.get('normalize', False)
    validate = param.get('validate', None)
    accept = param.get('accept', None)
    required = param.get('required', False)
    default = param.get('default', None)
    
    # Use attack values if provided, otherwise fall back to CLI args, then config defaults
    if attack_values and name in attack_values:
        value = attack_values[name]
    else:
        # Map param name to CLI arg, fall back to config default
        if name == 'email':
            value = USER_EMAIL or default
        elif name == 'file':
            value = USER_FILE or default
        elif name == 'url':
            value = USER_URL or default
        elif name == 'country':
            value = USER_COUNTRY or default
        else:
            # For other params, fallback to SAMPLE_TEXT, then default
            value = SAMPLE_TEXT or default
    
    # Apply normalization and validation
    if name == 'email':
        if normalize and not (attack_values and name in attack_values):
            # Only normalize/validate if not using attack values
            value = normalize_email(value)
            if required and not value:
                print(f"Error: Valid email required for parameter '{name}'")
                return None
        elif attack_values and name in attack_values:
            # For attack simulations, use the raw attack prompt without validation
            value = attack_values[name]
    elif name == 'file' and param_type == 'file':
        if attack_values and name in attack_values:
            # Generate malicious file with attack payload only when file type is required
            malicious_text = attack_values[name]
            if accept:
                if '.docx' in accept:
                    value = generate_malicious_docx(malicious_text)
                elif '.pdf' in accept:
                    value = generate_malicious_pdf(malicious_text)
                elif '.json' in accept:
                    value = generate_malicious_json(malicious_text)
                elif '.mp3' in accept:
                    value = generate_adversarial_audio(malicious_text)
                else:
                    # Default to creating a text file with the extension
                    ext = accept[0] if accept else '.txt'
                    temp_file = tempfile.NamedTemporaryFile(suffix=ext, delete=False, mode='w')
                    temp_file.write(f"Malicious content: {malicious_text}")
                    temp_file.close()
                    value = temp_file.name
            else:
                # No specific format required, create a generic file
                temp_file = tempfile.NamedTemporaryFile(delete=False, mode='w')
                temp_file.write(f"Malicious content: {malicious_text}")
                temp_file.close()
                value = temp_file.name
        elif accept and value and not (attack_values and name in attack_values):
            if not any(value.endswith(ext) for ext in accept):
                print(f"Error: File type for '{name}' must be one of {accept}")
                return None
        if required and not value:
            print(f"Error: File required for parameter '{name}'")
            return None
    elif name == 'url':
        if validate == 'url' and not (attack_values and name in attack_values):
            value = validate_url(value)
            if required and not value:
                print(f"Error: Valid URL required for parameter '{name}'")
                return None
        elif attack_values and name in attack_values:
            # For attack simulations, use the raw attack prompt without validation
            value = attack_values[name]
    elif name == 'country' and required and not value:
        print(f"Error: Country required for parameter '{name}'")
        return None
    
    return value

def is_rate_limit_or_connection_error(exception):
    """Check if the exception indicates rate limiting or connection issues"""
    error_str = str(exception).lower()
    return any(keyword in error_str for keyword in [
        'connection aborted', 'connection timeout', 'remote end closed', 
        'rate limit', 'too many requests', 'service unavailable',
        'connection reset', 'timeout', 'connection error'
    ])

def process_dynamic_api(endpoint_cfg, attack_values=None, max_retries=3, base_delay=2):
    url = endpoint_cfg['endpoint']
    method = endpoint_cfg.get('method', 'POST').upper()
    params = endpoint_cfg.get('parameters', [])
    files = {}
    data = {}
    for param in params:
        value = get_param_value(param, attack_values)
        if value is None:
            print(f"Skipping endpoint '{url}' due to missing/invalid parameter '{param['name']}'")
            return None
        if param['type'] == 'file':
            files[param['name']] = open(value, 'rb')
        else:
            data[param['name']] = value
    
    log_file = f"data/logs/log_{endpoint_cfg['name']}.csv"
    
    for attempt in range(max_retries + 1):
        try:
            if attempt > 0:
                delay = base_delay * (2 ** (attempt - 1))  # Exponential backoff
                print(f"⏰ Waiting {delay} seconds before retry {attempt}/{max_retries}...")
                time.sleep(delay)
            
            print(f"→ Sending request to {url} with data: {data}")
            
            if method == 'POST':
                response = requests.post(url, data=data, files=files if files else None, timeout=30)
            else:
                response = requests.get(url, params=data, timeout=30)
            response.raise_for_status()
            result = response.json() if response.headers.get('Content-Type', '').startswith('application/json') else response.text
            print(f"✓ Response from {url}: {result}")
            # Log result
            result_data = {
                "timestamp": datetime.datetime.now(datetime.timezone.utc).isoformat(),
                "endpoint": url,
                "parameters": data,
                "result": result,
                "status": "success",
                "misbehavior_detected": "false",
                "misbehavior_reason": ""
            }
            log_result_to_file(result_data, log_file)
            print(f"✓ Results logged to {log_file}")
            return True
        except Exception as e:
            if attempt < max_retries and is_rate_limit_or_connection_error(e):
                print(f"⚠️ Connection/rate limit error (attempt {attempt + 1}/{max_retries + 1}): {e}")
                continue  # Retry with backoff
            else:
                print(f"✗ Error calling {url}: {e}")
                # Log error result
                error_data = {
                    "timestamp": datetime.datetime.now(datetime.timezone.utc).isoformat(),
                    "endpoint": url,
                    "parameters": data,
                    "error": str(e),
                    "status": "error",
                    "misbehavior_detected": "false",
                    "misbehavior_reason": ""
                }
                log_result_to_file(error_data, log_file)
                print(f"✓ Error logged to {log_file}")
                return False
        finally:
            # Always close files on each attempt to avoid resource leaks
            for f in files.values():
                if hasattr(f, 'close'):
                    f.close()
            # Reopen files for retry if needed
            if attempt < max_retries:
                files = {}
                for param in params:
                    if param['type'] == 'file':
                        value = get_param_value(param, attack_values)
                        if value:
                            files[param['name']] = open(value, 'rb')


if __name__ == "__main__":

    print(f"Running BAS pipeline for all API endpoints defined in {CONFIG_PATH}")
    with open(CONFIG_PATH, 'r') as f:
        api_configs = json.load(f)

    processed_endpoints = []
    # Discover all attack modules in attacks folder
    attacks_dir = os.path.join(os.path.dirname(__file__), '../attacks')
    attack_module_files = [f for f in os.listdir(attacks_dir) if f.endswith('.py') and not f.startswith('__')]
    attack_module_names = [os.path.splitext(f)[0] for f in attack_module_files]

    for endpoint_cfg in api_configs:
        print(f"\n=== Running API attack: {endpoint_cfg['name']} ===")
        # For each attack module, generate attack payloads and call API
        for module_name in attack_module_names:
            print(f"--- Using attack module: {module_name} ---")
            attack_module = importlib.import_module(f'attacks.{module_name}')
            # Find all functions starting with 'simulate_' or 'generate_injected_prompt'
            attack_funcs = [getattr(attack_module, fn) for fn in dir(attack_module)
                            if fn.startswith('simulate_') or fn == 'generate_injected_prompt']
            for attack_func in attack_funcs:
                try:
                    prompts = attack_func(SAMPLE_TEXT) if SAMPLE_TEXT else attack_func('test')
                except TypeError:
                    prompts = [attack_func(SAMPLE_TEXT) if SAMPLE_TEXT else attack_func('test')]
                for prompt in prompts:
                    print(f"Testing prompt: {prompt}")
                    # Create attack values dictionary for this prompt
                    attack_values = {}
                    for param in endpoint_cfg.get('parameters', []):
                        # Only inject attack payloads into file parameters when type is "file"
                        # For other string parameters, use config defaults
                        if param['name'] == 'file' and param.get('type') == 'file':
                            attack_values[param['name']] = prompt
                        elif param['name'] in ['input', 'email', 'url'] and param.get('type') == 'string' and not param.get('default'):
                            # Only inject into string params that don't have defaults
                            attack_values[param['name']] = prompt
                    
                    result = process_dynamic_api(endpoint_cfg, attack_values)
                    # Only add to processed list if not skipped
                    if result is not None:
                        processed_endpoints.append(endpoint_cfg)
                    
                    # Small delay between requests to avoid overwhelming the API
                    time.sleep(0.5)

    # Step 2: Validate all generated logs
    print("\nValidating all generated logs...")
    for endpoint_cfg in processed_endpoints:
        log_file = f"data/logs/log_{endpoint_cfg['name']}.csv"
        if os.path.exists(log_file):
            os.system(f"python validation/run_validation.py {log_file}")
        else:
            print(f"Skipping validation for {log_file} (file not found)")
    # Step 3: Generate visual graph from all logs
    print("\nGenerating visual graph from logs...")
    log_files_str = ' '.join([f"data/logs/log_{cfg['name']}.csv" for cfg in processed_endpoints if os.path.exists(f"data/logs/log_{cfg['name']}.csv")])
    if log_files_str:
        os.system(f"python validation/report_generator.py {log_files_str} data/logs/report.html")
        print("✓ Visual graph generated at data/logs/report.html")
    else:
        print("No log files found for reporting.")